from io import BytesIO
from docx import Document
from django.core.files.base import ContentFile
from django.db import models

def main(action, request):
    if action == 'Upload & Replace' and 'document' in request.FILES:
        uploaded_doc = handle_uploaded_document(request.FILES['document'], request.POST.get('old_word', ''), request.POST.get('new_word', ''))
        return uploaded_doc

    elif action == 'Export':
        uploaded_doc = export_modified_document()
        return uploaded_doc

    elif action == 'jj_no_fish':
        upload_completed = False
        return {'upload_completed': upload_completed, 'old_word': '', 'new_word': ''}

    return {}

def handle_uploaded_document(file, old_word, new_word):
    new_document = UploadedDocument(document=file)
    new_document.save()
    modified_doc = replace_words_in_document(new_document.document, old_word, new_word)
    new_document.modified_document.save(f'{file.name}_modified.docx', ContentFile(modified_doc))
    return {
        'upload_completed': True,
        'old_word': old_word,
        'new_word': new_word,
        'modified_doc': modified_doc,
        'new_document_id': new_document.id  # Pass the ID for referencing the document later
    }


def replace_words_in_document(document_file, old_word, new_word):
    doc = Document(document_file)
    for paragraph in doc.paragraphs:
        if old_word in paragraph.text:
            paragraph.text = paragraph.text.replace(old_word, new_word)

    modified_document = BytesIO()
    doc.save(modified_document)
    modified_document.seek(0)
    print("replacement working")
    return modified_document.getvalue()

def export_modified_document():
    uploaded_doc = UploadedDocument.objects.last()
    if uploaded_doc and uploaded_doc.modified_document:
        file_name = uploaded_doc.modified_document.name
        file_name = file_name.replace("modified_documents_","")
        return {'modified_doc': uploaded_doc.modified_document.read(), 'new_file_name': file_name}
    else:
        return {}
