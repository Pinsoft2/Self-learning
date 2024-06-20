from io import BytesIO
import pytest
from docx import Document
from project import (
    handle_document,
    export_modified_document,
    upload_document,
    get_replacement_words,
)


@pytest.fixture
def sample_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("This is a sample document.")
    doc.add_paragraph("It contains the word 'test'.")
    file_path = tmp_path / "sample.docx"
    doc.save(file_path)
    return file_path

# Write test cases

def test_handle_document_with_regex_replacement(sample_docx):
    old_word = 'test'
    new_word = 'trial'

    # Create a BytesIO object for the sample document
    with open(sample_docx, 'rb') as file:
        content = file.read()
        modified_doc = BytesIO(content)

    # Replace the word using handle_document function
    modified_doc = handle_document(sample_docx, old_word, new_word)

    # Read the modified document
    doc = Document(modified_doc)
    paragraphs = [p.text for p in doc.paragraphs]

    # Assert that the word was replaced correctly
    assert 'test' not in paragraphs[1]  # Second paragraph should not contain the old word
    assert 'trial' in paragraphs[1]     # Second paragraph should contain the new word
    

def test_export_modified_document():
    # Prepare a sample modified document in memory
    modified_content = b"Modified document content"
    modified_doc = BytesIO()
    modified_doc.write(modified_content)
    modified_doc.seek(0)

    export_path = "test_export.docx"

    exported_path = export_modified_document(modified_doc, export_path)

    assert exported_path == export_path

def test_upload_document(monkeypatch):
    # Use monkeypatch to mock user input
    sample_file_path = "sample_document.docx"
    mock_input = lambda _: sample_file_path

    # Pass monkeypatch fixture to the function
    monkeypatch.setattr('builtins.input', mock_input)

    assert upload_document() == sample_file_path

def test_get_replacement_words(monkeypatch):
    # Use monkeypatch to mock user input
    old_word = "old_word"
    new_word = "new_word"
    mock_input = lambda _: old_word if "replace" in _ else new_word

    # Pass monkeypatch fixture to the function
    monkeypatch.setattr('builtins.input', mock_input)

    assert get_replacement_words() == (old_word, new_word)
