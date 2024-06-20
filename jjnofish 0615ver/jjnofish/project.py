from io import BytesIO
from docx import Document
from django.core.files.base import ContentFile
from django.db import models
import json
from .models import UploadedDocument

# project.py
def main(action, request, user=None, word_pairs=None):
    if action == 'Upload & Replace' and 'document' in request.FILES:
        uploaded_doc = handle_uploaded_document(request.FILES['document'],  word_pairs, user)
        return uploaded_doc

    elif action == 'Export':
        uploaded_doc = export_modified_document()
        return uploaded_doc

    elif action == 'jj_no_fish':
        upload_completed = False
        return {'upload_completed': upload_completed, 'old_words': [], 'new_words': []}

    return {}

def handle_uploaded_document(file, user, word_pairs):
    new_document = UploadedDocument(document=file)
    new_document.save()
    modified_doc = replace_words_in_document(new_document.document, word_pairs)
    new_document.modified_document.save(f'{file.name}_modified.docx', ContentFile(modified_doc))
    new_document.word_pairs = json.dumps(word_pairs)  # Save word pairs as JSON
    new_document.user = user
    new_document.save()
    return {
        'upload_completed': True,
        'word_pairs': word_pairs,
        'modified_doc': modified_doc,
        'new_document_id': new_document.id
    }


def replace_words_in_document(document_file, word_pairs):
    doc = Document(document_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for old_word, new_word in word_pairs.items():
            text = text.replace(old_word, new_word)
        paragraph.text = text

    modified_document = BytesIO()
    doc.save(modified_document)
    modified_document.seek(0)
    print("replacement working")
    return modified_document.getvalue()

def export_modified_document():
    uploaded_doc = UploadedDocument.objects.last()
    if uploaded_doc and uploaded_doc.modified_document:
        file_name = uploaded_doc.modified_document.name
        file_name = file_name.replace("modified_document_","")
        return {'modified_doc': uploaded_doc.modified_document.read(), 'new_file_name': file_name}
    else:
        return {}
